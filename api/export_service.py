"""
Export Service Module for AI Interviewer Platform
Provides export functionality for articles in PDF, Markdown, and DOCX formats
"""

import io
from datetime import datetime
from typing import Optional
import markdown2
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT

# Import removed - not needed for export service
from db.models import Interview, Article, User


class ExportService:
    """Service class for exporting articles in various formats"""
    
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_pdf_styles()
    
    def _setup_pdf_styles(self):
        """Setup custom PDF styles"""
        # Title style
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Title'],
            fontSize=24,
            textColor=colors.HexColor('#2E86AB'),
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        ))
        
        # Subtitle style
        self.styles.add(ParagraphStyle(
            name='CustomSubtitle',
            parent=self.styles['Normal'],
            fontSize=14,
            textColor=colors.HexColor('#666666'),
            spaceAfter=20,
            alignment=TA_CENTER,
            fontName='Helvetica-Oblique'
        ))
        
        # Metadata style
        self.styles.add(ParagraphStyle(
            name='Metadata',
            parent=self.styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#888888'),
            spaceAfter=15,
            fontName='Helvetica'
        ))
        
        # Enhanced body style
        self.styles.add(ParagraphStyle(
            name='CustomBody',
            parent=self.styles['Normal'],
            fontSize=12,
            textColor=colors.black,
            spaceAfter=12,
            alignment=TA_JUSTIFY,
            fontName='Helvetica',
            leading=16
        ))

    def export_to_pdf(self, article: Article, interview: Interview, user: Optional[User] = None) -> io.BytesIO:
        """
        Export article to PDF format with professional formatting
        
        Args:
            article: Article object to export
            interview: Associated interview object
            user: User who created the article (optional)
            
        Returns:
            io.BytesIO: PDF file buffer
        """
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        story = []
        
        # Title
        story.append(Paragraph(article.title, self.styles['CustomTitle']))
        story.append(Spacer(1, 20))
        
        # Subtitle with interview topic
        subtitle = f"Interview on: {interview.topic}"
        story.append(Paragraph(subtitle, self.styles['CustomSubtitle']))
        story.append(Spacer(1, 15))
        
        # Metadata section
        metadata_lines = [
            f"<b>Target Audience:</b> {interview.target_audience}",
            f"<b>Interview Date:</b> {interview.created_at.strftime('%B %d, %Y')}",
            f"<b>Article Version:</b> {article.version}",
            f"<b>Editor Iterations:</b> {article.editor_iterations}",
            f"<b>Generated:</b> {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
        ]
        
        if user:
            metadata_lines.insert(-1, f"<b>Author:</b> {user.username}")
        
        for line in metadata_lines:
            story.append(Paragraph(line, self.styles['Metadata']))
        
        story.append(Spacer(1, 30))
        
        # Article content with paragraph handling
        content_paragraphs = article.content.split('\n\n')
        for para in content_paragraphs:
            if para.strip():
                story.append(Paragraph(para.strip(), self.styles['CustomBody']))
                story.append(Spacer(1, 8))
        
        # Footer
        story.append(Spacer(1, 40))
        footer_text = "Generated by AI Interviewer Platform"
        story.append(Paragraph(footer_text, self.styles['Metadata']))
        
        doc.build(story)
        buffer.seek(0)
        return buffer

    def export_to_docx(self, article: Article, interview: Interview, user: Optional[User] = None) -> io.BytesIO:
        """
        Export article to DOCX format with professional formatting
        
        Args:
            article: Article object to export
            interview: Associated interview object
            user: User who created the article (optional)
            
        Returns:
            io.BytesIO: DOCX file buffer
        """
        doc = Document()
        
        # Configure document styles
        styles = doc.styles
        
        # Title
        title = doc.add_heading(article.title, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph(f"Interview on: {interview.topic}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_format = subtitle.runs[0].font
        subtitle_format.italic = True
        subtitle_format.size = Pt(14)
        
        doc.add_paragraph()  # Empty line
        
        # Metadata table
        table = doc.add_table(rows=5 + (1 if user else 0), cols=2)
        table.style = 'Light Shading Accent 1'
        
        metadata_items = [
            ("Target Audience", interview.target_audience),
            ("Interview Date", interview.created_at.strftime('%B %d, %Y')),
            ("Article Version", str(article.version)),
            ("Editor Iterations", str(article.editor_iterations)),
            ("Generated", datetime.now().strftime('%B %d, %Y at %I:%M %p'))
        ]
        
        if user:
            metadata_items.insert(-1, ("Author", user.username))
        
        for i, (label, value) in enumerate(metadata_items):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            
            # Format label cells
            label_run = row.cells[0].paragraphs[0].runs[0]
            label_run.font.bold = True
        
        doc.add_paragraph()  # Empty line
        
        # Article content
        content_paragraphs = article.content.split('\n\n')
        for para in content_paragraphs:
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph("Generated by AI Interviewer Platform")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_format = footer.runs[0].font
        footer_format.italic = True
        footer_format.size = Pt(10)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def export_to_markdown(self, article: Article, interview: Interview, user: Optional[User] = None) -> str:
        """
        Export article to Markdown format
        
        Args:
            article: Article object to export
            interview: Associated interview object
            user: User who created the article (optional)
            
        Returns:
            str: Markdown formatted content
        """
        markdown_content = []
        
        # Title
        markdown_content.append(f"# {article.title}")
        markdown_content.append("")
        
        # Subtitle
        markdown_content.append(f"*Interview on: {interview.topic}*")
        markdown_content.append("")
        
        # Metadata
        markdown_content.append("## Article Information")
        markdown_content.append("")
        markdown_content.append("| Field | Value |")
        markdown_content.append("|-------|--------|")
        markdown_content.append(f"| Target Audience | {interview.target_audience} |")
        markdown_content.append(f"| Interview Date | {interview.created_at.strftime('%B %d, %Y')} |")
        if user:
            markdown_content.append(f"| Author | {user.username} |")
        markdown_content.append(f"| Article Version | {article.version} |")
        markdown_content.append(f"| Editor Iterations | {article.editor_iterations} |")
        markdown_content.append(f"| Generated | {datetime.now().strftime('%B %d, %Y at %I:%M %p')} |")
        markdown_content.append("")
        
        # Article content
        markdown_content.append("## Article Content")
        markdown_content.append("")
        
        # Process content paragraphs
        content_paragraphs = article.content.split('\n\n')
        for para in content_paragraphs:
            if para.strip():
                markdown_content.append(para.strip())
                markdown_content.append("")
        
        # Footer
        markdown_content.append("---")
        markdown_content.append("*Generated by AI Interviewer Platform*")
        
        return "\n".join(markdown_content)

    def export_to_html(self, article: Article, interview: Interview, user: Optional[User] = None) -> str:
        """
        Export article to HTML format (bonus format)
        
        Args:
            article: Article object to export
            interview: Associated interview object
            user: User who created the article (optional)
            
        Returns:
            str: HTML formatted content
        """
        # First convert to markdown, then to HTML
        markdown_content = self.export_to_markdown(article, interview, user)
        
        # Convert markdown to HTML
        html_content = markdown2.markdown(
            markdown_content,
            extras=['tables', 'fenced-code-blocks', 'metadata']
        )
        
        # Wrap in a complete HTML document with CSS
        full_html = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{article.title}</title>
    <style>
        body {{
            font-family: 'Georgia', serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 40px auto;
            padding: 20px;
            color: #333;
            background-color: #fff;
        }}
        h1 {{
            color: #2E86AB;
            text-align: center;
            border-bottom: 3px solid #2E86AB;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #2E86AB;
            margin-top: 30px;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 12px;
            text-align: left;
        }}
        th {{
            background-color: #f8f9fa;
            font-weight: bold;
        }}
        .metadata {{
            background-color: #f8f9fa;
            padding: 15px;
            border-radius: 5px;
            margin: 20px 0;
        }}
        .footer {{
            text-align: center;
            margin-top: 40px;
            font-style: italic;
            color: #666;
            border-top: 1px solid #eee;
            padding-top: 20px;
        }}
    </style>
</head>
<body>
    {html_content}
</body>
</html>
        """
        
        return full_html

    def get_filename(self, article: Article, format_type: str) -> str:
        """
        Generate appropriate filename for the export
        
        Args:
            article: Article object
            format_type: Export format (pdf, docx, markdown, html)
            
        Returns:
            str: Generated filename
        """
        # Sanitize title for filename
        safe_title = "".join(c for c in article.title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        safe_title = safe_title.replace(' ', '_')[:50]  # Limit length
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        return f"article_{safe_title}_{timestamp}.{format_type}"


# Global export service instance
export_service = ExportService() 